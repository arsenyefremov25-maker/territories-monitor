from __future__ import annotations

from io import BytesIO

from docx import Document

from src.config import CATEGORY_ACTIVE, CATEGORY_OCCUPIED, CATEGORY_POSSIBLE
from src.parser import extract_rows, make_record_key, normalize_date


def build_current_layout_docx() -> bytes:
    document = Document()
    table = document.add_table(rows=1, cols=9)
    headers = [
        "Область",
        "Район",
        "Територіальна громада",
        "Населений пункт",
        "Код",
        "Категорія",
        "Дата початку",
        "Дата завершення",
        "Функціонування систем",
    ]
    for index, value in enumerate(headers):
        table.rows[0].cells[index].text = value

    rows = [
        (
            "Дніпропетровська",
            "Нікопольський",
            "Тестова міська територіальна громада",
            "м. Тестове",
            "UA12080010010085669",
            "Територія активних бойових дій",
            "01.06.2023",
            "13.10.2023",
            "Так",
        ),
        (
            "Дніпропетровська",
            "Нікопольський",
            "Тестова міська територіальна громада",
            "м. Тестове",
            "UA12080010010085669",
            "Територія можливих бойових дій",
            "14.10.2023",
            "",
            "Так",
        ),
        (
            "Херсонська",
            "Генічеський",
            "Інша селищна територіальна громада",
            "с-ще Інше",
            "UA65040010010012345",
            "Тимчасово окупована Російською Федерацією територія України",
            "24.02.2022",
            "",
            "Ні",
        ),
        (
            "Сумська",
            "Охтирський",
            "Тростянецька міська територіальна громада",
            "с. Новоселівка",
            "UA59040130230084876",
            "Тимчасово окупована Російською Федерацією територія України",
            "24.02.2022",
            "26.03.2022",
            "Ні",
        ),
        (
            "Сумська",
            "Охтирський",
            "Тростянецька міська територіальна громада",
            "с. Новоукраїнка",
            "UA59040130230084876",
            "Тимчасово окупована Російською Федерацією територія України",
            "24.02.2022",
            "26.03.2022",
            "Ні",
        ),
    ]

    for values in rows:
        cells = table.add_row().cells
        for index, value in enumerate(values):
            cells[index].text = value

    buffer = BytesIO()
    document.save(buffer)
    return buffer.getvalue()


def build_legacy_layout_docx() -> bytes:
    document = Document()
    document.add_paragraph("1. Території можливих бойових дій")
    table = document.add_table(rows=4, cols=4)
    table.rows[0].cells[0].text = "ДОНЕЦЬКА ОБЛАСТЬ"
    table.rows[1].cells[0].text = "Краматорський район"
    table.rows[2].cells[0].text = "Код"
    table.rows[2].cells[1].text = "Найменування"
    table.rows[2].cells[2].text = "Дата початку"
    table.rows[2].cells[3].text = "Дата завершення"
    table.rows[3].cells[0].text = "UA14120010010000123"
    table.rows[3].cells[1].text = "Перша громада"
    table.rows[3].cells[2].text = "01.01.2025"
    table.rows[3].cells[3].text = ""

    buffer = BytesIO()
    document.save(buffer)
    return buffer.getvalue()


def test_extract_rows_parses_current_nine_column_layout() -> None:
    frame = extract_rows(build_current_layout_docx())
    assert len(frame) == 5
    assert set(frame["category"]) == {
        CATEGORY_POSSIBLE,
        CATEGORY_ACTIVE,
        CATEGORY_OCCUPIED,
    }
    assert set(frame["systems_functioning"].dropna()) == {True, False}
    assert frame["hromada_name"].notna().all()
    assert frame["settlement_name"].notna().all()
    assert frame["record_key"].is_unique


def test_extract_rows_keeps_legacy_layout_compatibility() -> None:
    frame = extract_rows(build_legacy_layout_docx())
    assert len(frame) == 1
    assert frame.iloc[0]["category"] == CATEGORY_POSSIBLE
    assert frame.iloc[0]["hromada_name"] == "Перша громада"


def test_record_key_is_stable_for_same_identity() -> None:
    start = normalize_date("01.01.2025")
    first = make_record_key(
        "UA14120010010000123",
        CATEGORY_ACTIVE,
        start,
        identity_hint="Донецька|Краматорський|Тестове",
    )
    second = make_record_key(
        "UA14120010010000123",
        CATEGORY_ACTIVE,
        start,
        identity_hint="Донецька|Краматорський|Тестове",
    )
    assert first == second
